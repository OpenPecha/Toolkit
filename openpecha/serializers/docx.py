from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from openpecha.serializers.epub import EpubSerializer
from openpecha.serializers.serialize import Serialize


class DocxSerializer(Serialize):

    def get_all_p_tags(self, html):
        """list all the p tags in serialized html

        Args:
            html (string): serialized html

        Returns:
            list: list of para tags in serialized html
        """
        soup = BeautifulSoup(html, "html.parser")
        p_tags = soup.find_all("p")
        return p_tags

    def format_span(self, span, para_doc_obj, document):
        """Add text and its custom style of span tag to document

        Args:
            span (bs4 obj): span tag
            para_doc_obj (para obj): para obj of document
            document (obj): document object in which text with its styles are accumilated
        """
        if span['class'][0] == "tibetan-root-text-in-verse":
            para_doc_obj.style = document.styles['Tsawa_verse_first_line']
            Tsawa = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-root-text-in-verse-middle-line" == span['class'][0]:
            para_doc_obj.style = document.styles['Tsawa_verse_middle_line']
            Tsawa = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-root-text-in-verse-last-line" == span['class'][0]:
            para_doc_obj.style = document.styles['Tsawa_verse_last_line']
            Tsawa = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-root-text" == span['class'][0]:
            Tsawa = para_doc_obj.add_run(f'{span.text}')
            Tsawa.style = document.styles['Tsawa_inline']

        elif span['class'][0] == "tibetan-citations-in-verse":
            para_doc_obj.style = document.styles['Citation_verse_first_line']
            Citation = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-citations-in-verse-middle-line" == span['class'][0]:
            para_doc_obj.style = document.styles['Citation_verse_middle_line']
            Citation = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-citations-in-verse-last-line" == span['class'][0]:
            para_doc_obj.style = document.styles['Citation_verse_last_line']
            Citation = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-citations-prose" == span['class'][0]:
            para_doc_obj.style = document.styles['Citation_prose']
            Citation = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-citations" == span['class'][0]:
            Citation = para_doc_obj.add_run(f'{span.text}')
            Citation.style = document.styles['Citation_inline']

        elif "tibetan-sabche" == span['class'][0]:
            para_doc_obj.style = document.styles['Heading 2']
            Sabche = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-sabche1" == span['class'][0]:
            Inline_sabche = para_doc_obj.add_run(f'{span.text}')
            Inline_sabche.style = document.styles['Inline_sabche']

        elif "tibetan-chapters" in span['class'][0]:
            para_doc_obj.style = document.styles['Heading 1']
            Chapter = para_doc_obj.add_run(f'{span.text}')

        elif "credits-page_front-title" in span['class'][0]:
            para_doc_obj.style = document.styles['Book_title']
            Booktitle = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-book-sub-title" in span['class'][0]:
            para_doc_obj.style = document.styles['Subtitle']
            Subtitle = para_doc_obj.add_run(f'{span.text}')

        elif "front-page---text-titles" in span['class'][0]:
            para_doc_obj.style = document.styles['Author']
            Author = para_doc_obj.add_run(f'{span.text}')

        elif "tibetan-commentary-small" == span['class'][0]:
            Yigchung = para_doc_obj.add_run(f'{span.text}')
            Yigchung.style = document.styles['Yigchung']

        else:
            normal_text = para_doc_obj.add_run(f'{span.text}')
            normal_text.style = document.styles['Commentary']

    def add_credit_page(self, p_tag, para_doc_obj):
        """Serialise credit page iamge in document

        Args:
            p_tag (obj): para tag object of bs4
            para_doc_obj (obj): para obj of document
        """
        img_tag = p_tag.find_all('img')
        img_path = img_tag[0]['src']
        credit_page = para_doc_obj.add_run()
        credit_page.add_picture(img_path)

    def page_break_needed(self, p_tag):
        """Return true of p tag's text need a page break after it

        Args:
            p_tag (obj): para tag object of bs4

        Returns:
            Boolean: True if page break needed after wards else False
        """
        p_tag_class = p_tag['class'][0]
        if p_tag_class == 'tibetan-commentary-non-indent1' or p_tag_class == 'tibetan-regular-indented1':
            return True
        else:
            return False

    def format_p_tag(self, p_tag, para_doc_obj, document):
        """Format p tag text with its custom style in document

        Args:
            p_tag (obj): para tag object of bs4
            para_doc_obj (obj): para obj of document
            document (obj): document object
        """
        if 'tibetan-regular-indented' in p_tag['class'][0]:
                para_doc_obj.style = document.styles['Indented_commentary']
        for content in p_tag.contents:
            if 'span' == content.name:
                self.format_span(content, para_doc_obj, document)
            else:
                normal_text = para_doc_obj.add_run(f'{content}')
                normal_text.style = document.styles['Commentary']
                if self.page_break_needed(p_tag):
                    normal_text.add_break = WD_BREAK.PAGE

    def add_styles(self, doc_styles):
        """Add all the custom style in document's style

        Args:
            doc_styles (obj): document's styles
        """

        #Book Title
        book_title_style = doc_styles.add_style('Book_title', WD_STYLE_TYPE.PARAGRAPH)
        book_title_font = book_title_style.font
        book_title_font.color.rgb = RGBColor(0,0,0)
        book_title_font.name = "Monlam Uni OuChan2"
        book_title_font.bold = True
        book_title_font.size = Pt(20)
        title_para_format = book_title_style.paragraph_format
        title_para_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_para_format.space_after = Pt(30)
        title_para_format.space_before = Pt(15)

        #Subtitle
        subtitle_style = doc_styles['Subtitle']
        subtitle_style.font.color.rgb = RGBColor(0,0,0)
        subtitle_style.font.italic = False
        subtitle_style.font.name = "Monlam Uni OuChan2"
        subtitle_para_format = subtitle_style.paragraph_format
        subtitle_para_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_para_format.space_after = Pt(14)
        subtitle_para_format.space_before = Pt(14)

        #Author
        author_style = doc_styles.add_style('Author', WD_STYLE_TYPE.PARAGRAPH)
        author_style.font.color.rgb = RGBColor(0,0,0)
        author_style.font.size = Pt(10)
        author_style.font.name = "Monlam Uni OuChan2"
        author_para_format = author_style.paragraph_format
        author_para_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        author_para_format.space_after = Pt(18)

        #Chapter
        chapter_style = doc_styles['Heading 1']
        chapter_style.font.color.rgb = RGBColor(0,0,0)
        chapter_style.font.name = "Monlam Uni OuChan2"
        chapter_style.font.bold = False
        chapter_style.font.size = Pt(16)
        chapter_para_format = chapter_style.paragraph_format
        chapter_para_format.space_after = Pt(14)
        chapter_para_format.space_before = Pt(14)


        #Sabche
        sabche_style = doc_styles['Heading 2']
        sabche_style.font.color.rgb = RGBColor(0,94,127)
        sabche_style.font.bold = False
        sabche_style.font.size = Pt(14)
        sabche_style.font.name = "Monlam Uni OuChan2"
        sabche_para_format = sabche_style.paragraph_format
        sabche_para_format.space_after = Pt(23)
        sabche_para_format.space_before = Pt(5)

        sabche_inline_style = doc_styles.add_style('Inline_sabche', WD_STYLE_TYPE.CHARACTER)
        sabche_inline_style.font.color.rgb = RGBColor(0,94,127)
        sabche_inline_style.font.size = Pt(14)
        sabche_inline_style.font.name = "Monlam Uni OuChan2"

        #Tsawa
        tsawa_style = doc_styles.add_style('Tsawa_verse_first_line', WD_STYLE_TYPE.PARAGRAPH)
        tsawa_style.font.color.rgb = RGBColor(139,20,9)
        tsawa_style.font.size = Pt(14)
        tsawa_style.font.name = "Monlam Uni OuChan2"
        tsawa_para_format = tsawa_style.paragraph_format
        tsawa_para_format.left_indent = Pt(48)
        tsawa_para_format.space_after = Pt(0)
        tsawa_para_format.space_before = Pt(5)

        tsawa_style = doc_styles.add_style('Tsawa_verse_middle_line', WD_STYLE_TYPE.PARAGRAPH)
        tsawa_style.font.color.rgb = RGBColor(139,20,9)
        tsawa_style.font.size = Pt(14)
        tsawa_style.font.name = "Monlam Uni OuChan2"
        tsawa_para_format = tsawa_style.paragraph_format
        tsawa_para_format.left_indent = Pt(48)
        tsawa_para_format.space_after = Pt(0)
        tsawa_para_format.space_before = Pt(0)

        tsawa_style = doc_styles.add_style('Tsawa_verse_last_line', WD_STYLE_TYPE.PARAGRAPH)
        tsawa_style.font.color.rgb = RGBColor(139,20,9)
        tsawa_style.font.size = Pt(14)
        tsawa_style.font.name = "Monlam Uni OuChan2"
        tsawa_para_format = tsawa_style.paragraph_format
        tsawa_para_format.left_indent = Pt(48)
        tsawa_para_format.space_after = Pt(20)
        tsawa_para_format.space_before = Pt(0)

        tsawa_inline_style = doc_styles.add_style('Tsawa_inline', WD_STYLE_TYPE.CHARACTER)
        tsawa_inline_style.font.color.rgb = RGBColor(139,20,9)
        tsawa_inline_style.font.size = Pt(14)
        tsawa_inline_style.font.name = "Monlam Uni OuChan2"

        #Citation
        citation_style = doc_styles.add_style('Citation_verse_first_line', WD_STYLE_TYPE.PARAGRAPH)
        citation_style.font.color.rgb = RGBColor(137,115,53)
        citation_style.font.size = Pt(14)
        citation_style.font.name = "Monlam Uni OuChan2"
        citation_para_format = citation_style.paragraph_format
        citation_para_format.left_indent = Pt(48)
        citation_para_format.space_after = Pt(0)
        citation_para_format.space_before = Pt(5)

        citation_style = doc_styles.add_style('Citation_verse_middle_line', WD_STYLE_TYPE.PARAGRAPH)
        citation_style.font.color.rgb = RGBColor(137,115,53)
        citation_style.font.size = Pt(14)
        citation_style.font.name = "Monlam Uni OuChan2"
        citation_para_format = citation_style.paragraph_format
        citation_para_format.left_indent = Pt(48)
        citation_para_format.space_after = Pt(0)
        citation_para_format.space_before = Pt(0)

        citation_style = doc_styles.add_style('Citation_verse_last_line', WD_STYLE_TYPE.PARAGRAPH)
        citation_style.font.color.rgb = RGBColor(137,115,53)
        citation_style.font.size = Pt(14)
        citation_style.font.name = "Monlam Uni OuChan2"
        citation_para_format = citation_style.paragraph_format
        citation_para_format.left_indent = Pt(48)
        citation_para_format.space_after = Pt(20)
        citation_para_format.space_before = Pt(0)

        citation_style = doc_styles.add_style('Citation_prose', WD_STYLE_TYPE.PARAGRAPH)
        citation_style.font.color.rgb = RGBColor(137,115,53)
        citation_style.font.size = Pt(14)
        citation_style.font.name = "Monlam Uni OuChan2"
        citation_para_format = citation_style.paragraph_format
        citation_para_format.left_indent = Pt(48)
        citation_para_format.right_indent = Pt(48)
        citation_para_format.space_after = Pt(20)
        citation_para_format.space_before = Pt(20)

        citation_inline_style = doc_styles.add_style('Citation_inline', WD_STYLE_TYPE.CHARACTER)
        citation_inline_style.font.color.rgb = RGBColor(137,115,53)
        citation_inline_style.font.size = Pt(14)
        citation_inline_style.font.name = "Monlam Uni OuChan2"

        #Yigchung
        yigchung_style = doc_styles.add_style('Yigchung', WD_STYLE_TYPE.CHARACTER)
        yigchung_style.font.color.rgb = RGBColor(0,0,0)
        yigchung_style.font.size = Pt(10)
        yigchung_style.font.name = "Monlam Uni OuChan2"

        #Commentary
        commentary_indent1_style = doc_styles.add_style('Indented_commentary', WD_STYLE_TYPE.PARAGRAPH)
        commentary_indent1_style.font.color.rgb = RGBColor(0,0,0)
        commentary_indent1_style.font.size = Pt(14)
        commentary_indent1_style.font.name = "Monlam Uni OuChan2"
        commentary_indent1_para_format = commentary_indent1_style.paragraph_format
        commentary_indent1_para_format.first_line_indent = Pt(18)

        commentary_style = doc_styles.add_style('Commentary', WD_STYLE_TYPE.CHARACTER)
        commentary_style.font.color.rgb = RGBColor(0,0,0)
        commentary_style.font.size = Pt(14)
        commentary_style.font.name = "Monlam Uni OuChan2"

    def create_docx(self, html, output_path, pecha_id):
        """Serialize html to docx file with custom style for annotated tags

        Args:
            html (string): serialized html using epubserializer
            output_path (obj): output path where docx file need to be saved
            pecha_id (string): pecha id
        """
        p_tags = self.get_all_p_tags(html)
        document = Document()
        doc_styles = document.styles
        self.add_styles(doc_styles)
        for p_tag in p_tags:
            para_doc_obj = document.add_paragraph()
            if p_tag.attrs and p_tag['class'][0] == 'credits-page_epub-edition-line':
                self.add_credit_page(p_tag, para_doc_obj)
            else:
                self.format_p_tag(p_tag, para_doc_obj, document)
        out_fn = output_path / f"{pecha_id}.docx"
        document.save(str(out_fn))
        return out_fn

    def serialize(self, output_path: Path, **kwargs) -> Path:
        output_path = Path(output_path)
        epub_serializer = EpubSerializer(self.opf_path)
        epub_serializer.apply_layers()
        pecha_title = epub_serializer.meta["source_metadata"].get("title", "")
        pecha_id = epub_serializer.meta["id"]
        results = epub_serializer.get_result()
        for base_id, result in results.items():
            serialized_html = epub_serializer.get_serialized_html(result, base_id, pecha_title)
            out_fn = self.create_docx(serialized_html, output_path, pecha_id)
            return out_fn
